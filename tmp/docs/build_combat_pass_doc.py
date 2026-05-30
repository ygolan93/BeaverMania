"""Generate Combat Visual Stability Pass DOCX."""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

OUT = Path(__file__).resolve().parents[2] / "output" / "doc" / "BeaverMania_Combat_Visual_Stability_Pass.docx"


def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def add_bullet(doc, text, level=0):
    style = "List Bullet" if level == 0 else "List Bullet 2"
    doc.add_paragraph(text, style=style)


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("BeaverMania - Combat Visual Stability Pass")
    run.bold = True
    run.font.size = Pt(18)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run("Implementation summary and Play Mode verification guide").italic = True

    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.add_run("Branch: ").bold = True
    meta.add_run("cursor/6838e3e6 (commit ba8344bb)\n")
    meta.add_run("Project: ").bold = True
    meta.add_run("BeaverMania Unity 3D action/platformer\n")
    meta.add_run("Date: ").bold = True
    meta.add_run("May 2026")

    add_heading(doc, "1. Overview", 1)
    doc.add_paragraph(
        "This pass improves local stability, combat readability, and visual polish without "
        "reducing enemy swarm density. Pause/menu behavior was regression-tested only; no "
        "pause code changes were required."
    )

    add_heading(doc, "2. Tasks delivered", 1)

    tasks = [
        (
            "TASK 01 - Footstep dirt/dust material (P0)",
            [
                "Root cause: procedural footstep particles had no ParticleSystemRenderer material (magenta under URP).",
                "Added M_FootstepDust_Cartoon.mat and material assignment in FootstepVfxEmitter.",
                "Removed temporary agent debug logging from footstep code.",
            ],
        ),
        (
            "TASK 02 - Boost gating behind combat (P0)",
            [
                "Boost (Goblet, Y key) now requires combat-earned charge, not free inventory use.",
                "New BoostChargeController, BoostChargeSettings (ScriptableObject), BoostChargeHudPresenter.",
                "Hooks on NPC_Basic and ScorpionScript for hits and kills.",
            ],
        ),
        (
            "TASK 04/05 - Enemy rendering and readability (P0/P1)",
            [
                "EnemyHealthBarVisibility: hide distant enemy health bars.",
                "NPC_Basic: far-AI update throttling, timed HitEffect, SetMaxNPCHealth on spawn.",
                "Shorter combat speech duration (ChangeSpeech 0.55s).",
                "Enemy count and spawn rates unchanged.",
            ],
        ),
        (
            "TASK 03 - Out-of-bounds containment (P1)",
            [
                "LevelBoundsController: Y void and horizontal bounds, checkpoint respawn, GobletOFF on reset.",
                "Gizmos-only debug visualization in Scene view.",
                "Auto-added on GameMaster at runtime.",
            ],
        ),
        (
            "TASK 06 - Hide debug trigger meshes (P1)",
            [
                "TriggerVolumeVisualHider disables MeshRenderer/SpriteRenderer on trigger volumes.",
                "GameplayTriggerVisualBootstrap applies at scene load (excludes Life checkpoints).",
            ],
        ),
        (
            "TASK 07/08 - Objectives and bridge UX (P1/P2)",
            [
                "ObjectiveSyncService and GameProgressionStage for single-writer objective flow.",
                "Trader dialogue completion advances objectives; bridge proximity UI shows carried vs placed logs.",
            ],
        ),
        (
            "TASK 09 - Pause regression (P3)",
            [
                "Manual checklist only (Assets/Docs/PAUSE_REGRESSION_CHECKLIST.md).",
                "No code changes unless a reproducible bug appears.",
            ],
        ),
    ]

    for title_text, bullets in tasks:
        add_heading(doc, title_text, 2)
        for b in bullets:
            add_bullet(doc, b)

    add_heading(doc, "3. Key files added or modified", 1)
    files = [
        "Assets/Materials/VFX/M_FootstepDust_Cartoon.mat",
        "Assets/Scripts/Display/FootstepVfxEmitter.cs",
        "Assets/Scripts/Player/Combat/BoostChargeController.cs",
        "Assets/Scripts/Data/Combat/BoostChargeSettings.cs",
        "Assets/Scripts/UI/Hud/BoostChargeHudPresenter.cs",
        "Assets/Scripts/NPC/EnemyHealthBarVisibility.cs",
        "Assets/Scripts/NPC/Wasp/NPC_Basic.cs",
        "Assets/Scripts/Core/GameFlow/LevelBoundsController.cs",
        "Assets/Scripts/Core/GameFlow/ObjectiveSyncService.cs",
        "Assets/Scripts/Core/GameFlow/TriggerVolumeVisualHider.cs",
        "Assets/Scripts/Objects/Newbridge/NewConstructor.cs",
        "Assets/Scripts/NPC/Beaver/Trader.cs",
        "Assets/Prefabs/OtterPlayer/Otter_Shapekeys/Player.prefab",
    ]
    for f in files:
        add_bullet(doc, f)

    add_heading(doc, "4. Manual Unity inspector checklist", 1)
    checks = [
        "Player prefab: assign M_FootstepDust_Cartoon on FootstepVfxEmitter.",
        "Player: BoostChargeController present; optional BoostChargeSettings asset assigned.",
        "PlayerCanvas: wire BoostChargeHudPresenter fill/status text if using visual meter bar.",
        "GameMaster / LevelBoundsController: tune boundsCenter, boundsSize, minWorldY in Scene view.",
        "Wasp prefab: EnemyHealthBarVisibility on health canvas; slider max matches MaxHealth.",
        "WayPoint Objective[] strings align with GameProgressionStage order.",
        "Bridge: NewConstructor BridgeUI reference and trigger collider for proximity prompts.",
    ]
    for i, c in enumerate(checks, 1):
        doc.add_paragraph(f"{i}. {c}")

    add_heading(doc, "5. Play Mode validation", 1)
    validation = [
        "Footsteps: 2+ minutes movement on dirt/grass/bridge - no magenta particles; subtle dust only.",
        "Boost: at start Y does nothing; combat fills meter; at threshold Y activates boost once; meter empty after use.",
        "Combat swarm: distant wasps hide health bars; FPS stable or improved vs same test area.",
        "OOB: falling below terrain or leaving bounds respawns at checkpoint; no visible debug cubes in Game view.",
        "Objectives: trader dialogue clears trader objective; logs/bridge objectives advance in order.",
        "Bridge: proximity prompt shows carried vs placed logs; prompt hides when leaving range.",
        "Pause: open, resume, menu cursor, return to gameplay (see PAUSE_REGRESSION_CHECKLIST.md).",
    ]
    for v in validation:
        add_bullet(doc, v)

    add_heading(doc, "6. Known follow-ups", 1)
    add_bullet(doc, "Level 1 scene bounds may need designer tuning after playtesting boost traversal.")
    add_bullet(doc, "TipTriggerZone scene placements remain optional Editor follow-up.")
    add_bullet(doc, "BoostChargeSettings_Default.asset can be authored in Editor for tuning without code changes.")

    doc.add_paragraph()
    footer = doc.add_paragraph("Generated for BeaverMania project documentation.")
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
